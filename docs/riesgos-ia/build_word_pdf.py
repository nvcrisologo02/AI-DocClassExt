from pathlib import Path
import re
import subprocess
import tempfile
import os
import shutil

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Image as RLImage
from PIL import Image as PILImage


def parse_markdown_lines(md_path: Path, out_img_dir: Path = None):
    """
    Parse markdown into token list. If `out_img_dir` is provided, pre-render
    ```mermaid``` blocks to PNG files inside that directory (using npx mermaid-cli)
    and return tokens ('image', path) for the generated images.
    """
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    parsed = []
    i = 0
    mermaid_count = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        # detect mermaid code fence start
        if re.match(r"^\s*```mermaid\b", line):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not re.match(r"^\s*```\s*$", lines[i]):
                mermaid_lines.append(lines[i])
                i += 1
            # skip closing fence if present
            # generate image if requested
            mermaid_count += 1
            if out_img_dir is not None:
                out_img_dir.mkdir(parents=True, exist_ok=True)
                img_name = f"{md_path.stem}_mermaid_{mermaid_count}.png"
                out_file = (out_img_dir / img_name).resolve()
                # write temp input
                with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.mmd', encoding='utf-8') as tf:
                    tf.write('\n'.join(mermaid_lines))
                    temp_name = tf.name
                try:
                    npx_cmd = shutil.which('npx')
                    if npx_cmd:
                        cmd = [npx_cmd, '--yes', '@mermaid-js/mermaid-cli', '-i', temp_name, '-o', str(out_file), '-w', '1200']
                        try:
                            res = subprocess.run(cmd, check=False, capture_output=True, text=True)
                        except FileNotFoundError as ex:
                            print(f"Warning: failed to execute npx for {md_path}: {ex}")
                            parsed.append(('p', '\n'.join(mermaid_lines)))
                        else:
                            if res.returncode != 0:
                                print(f"Warning: mermaid-cli failed for {md_path} block {mermaid_count}: {res.stderr.strip()}")
                                parsed.append(('p', '[MERMAID_RENDER_FAILED]'))
                            else:
                                parsed.append(('image', str(out_file)))
                    else:
                        print("Warning: 'npx' not found. Mermaid blocks will not be rendered.")
                        parsed.append(('p', '\n'.join(mermaid_lines)))
                finally:
                    try:
                        os.unlink(temp_name)
                    except Exception:
                        pass
            else:
                parsed.append(('p', '\n'.join(mermaid_lines)))
            # advance past closing fence if any
            while i < len(lines) and re.match(r"^\s*```\s*$", lines[i]):
                i += 1
            continue

        # normal markdown lines
        if not line:
            parsed.append(("blank", ""))
        elif line.startswith("### "):
            parsed.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            parsed.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            parsed.append(("h1", line[2:].strip()))
        elif line.startswith("- "):
            parsed.append(("bullet", line[2:].strip()))
        elif len(line) >= 3 and line[:3].isdigit() and line[1:3] == ". ":
            parsed.append(("number", line.strip()))
        else:
            # image markdown: ![alt](path)
            m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
            if m:
                img_path = m.group(1)
                parsed.append(("image", img_path))
            else:
                parsed.append(("p", line))
        i += 1
    return parsed


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    out_img_dir = docx_path.parent / 'images'
    for kind, text in parse_markdown_lines(md_path, out_img_dir=out_img_dir):
        if kind == "blank":
            doc.add_paragraph("")
        elif kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        elif kind == "image":
            try:
                img_path = str(text)
                if not os.path.isabs(img_path):
                    img_path = str((out_img_dir / img_path).resolve()) if out_img_dir.exists() else img_path
                doc.add_picture(img_path, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"[IMAGE ERROR: {e}]")
        else:
            doc.add_paragraph(text)
    doc.save(docx_path)


def md_to_pdf(md_path: Path, pdf_path: Path):
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=16,
        leading=19,
        spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=13,
        leading=16,
        spaceAfter=8,
    )
    h3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontSize=11,
        leading=14,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontSize=10,
        leading=13,
        spaceAfter=5,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=12,
    )

    left_margin = 2 * cm
    right_margin = 2 * cm
    top_margin = 2 * cm
    bottom_margin = 2 * cm
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
    )

    story = []
    out_img_dir = pdf_path.parent / 'images'
    tokens = parse_markdown_lines(md_path, out_img_dir=out_img_dir)
    for kind, text in tokens:
        safe_text = (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        if kind == "blank":
            story.append(Spacer(1, 6))
        elif kind == "h1":
            story.append(Paragraph(safe_text, h1))
        elif kind == "h2":
            story.append(Paragraph(safe_text, h2))
        elif kind == "h3":
            story.append(Paragraph(safe_text, h3))
        elif kind == "bullet":
            story.append(Paragraph("• " + safe_text, bullet))
        elif kind == "number":
            story.append(Paragraph(safe_text, body))
        elif kind == "image":
            try:
                img_path = text
                # if relative path, resolve against out_img_dir if exists
                if not os.path.isabs(img_path):
                    candidate = (out_img_dir / img_path).resolve()
                    if candidate.exists():
                        img_path = str(candidate)
                # compute width and scale
                available_width = getattr(doc, 'width', A4[0] - left_margin - right_margin)
                pil = PILImage.open(img_path)
                w_px, h_px = pil.size
                ratio = (h_px / w_px) if w_px else 1
                desired_width = available_width
                desired_height = desired_width * ratio
                # ensure image fits within page height (avoid LayoutError)
                page_height_available = A4[1] - top_margin - bottom_margin
                # leave a small safety margin (avoid exact-fit layout errors)
                safety_margin = 24  # points
                max_height = max(0, page_height_available - safety_margin)
                if desired_height > max_height:
                    desired_height = max_height
                    desired_width = desired_height / ratio if ratio else desired_width
                story.append(RLImage(img_path, width=desired_width, height=desired_height))
            except Exception as e:
                story.append(Paragraph(f"[IMAGE ERROR: {e}]", body))
        else:
            story.append(Paragraph(safe_text, body))

    try:
        doc.build(story)
    except PermissionError as e:
        print(f"Warning: no se pudo escribir el PDF {pdf_path}: {e}. ¿Está abierto? Se omite.")
        return


def main():
    import argparse
    base = Path(__file__).resolve().parent

    default_docs = [
        (
            base / "Documento_Riesgos_IA_Completo.md",
            base / "Documento_Riesgos_IA_Completo.docx",
            base / "Documento_Riesgos_IA_Completo.pdf",
        ),
        (
            base / "Documento_Riesgos_IA_Resumen.md",
            base / "Documento_Riesgos_IA_Resumen.docx",
            base / "Documento_Riesgos_IA_Resumen.pdf",
        ),
    ]

    parser = argparse.ArgumentParser(description="Convert Markdown files to DOCX and PDF.")
    parser.add_argument('files', nargs='*', help='Markdown files to convert (paths). If omitted, uses defaults in script.')
    parser.add_argument('--only-pdf', action='store_true', help='Generate only PDF (skip DOCX).')
    parser.add_argument('--all-docs', action='store_true', help='Process all .md files under docs root.')
    parser.add_argument('--docs-root', default=str(Path.cwd() / 'docs'), help='Root folder to search for MD files when using --all-docs.')
    parser.add_argument('--out-root', default=None, help='Output root folder for generated PDFs (mirrors structure). Defaults to docs_root/pdf')
    args = parser.parse_args()

    docs = []
    if args.all_docs:
        docs_root = Path(args.docs_root).resolve()
        if not docs_root.exists():
            print(f"Error: docs root no existe: {docs_root}")
            return
        out_root = Path(args.out_root).resolve() if args.out_root else (docs_root / 'pdf')
        for md_path in docs_root.rglob('*.md'):
            # Skip files inside any generated pdf folder
            if 'pdf' in [p.lower() for p in md_path.parts]:
                continue
            rel = md_path.relative_to(docs_root)
            pdf_path = (out_root / rel).with_suffix('.pdf')
            docx_path = md_path.with_suffix('.docx')
            docs.append((md_path, docx_path, pdf_path))
    elif args.files:
        out_root = Path(args.out_root).resolve() if args.out_root else None
        docs_root = Path(args.docs_root).resolve() if args.docs_root else Path.cwd() / 'docs'
        for f in args.files:
            md_path = Path(f)
            if not md_path.is_absolute():
                md_path = (Path.cwd() / f).resolve()
            docx_path = md_path.with_suffix('.docx')
            if out_root:
                try:
                    rel = md_path.relative_to(docs_root)
                except Exception:
                    rel = Path(md_path.name)
                pdf_path = (out_root / rel).with_suffix('.pdf')
            else:
                pdf_path = md_path.with_suffix('.pdf')
            docs.append((md_path, docx_path, pdf_path))
    else:
        docs = default_docs

    generated = []
    for md_path, docx_path, pdf_path in docs:
        if not md_path.exists():
            print(f"Advertencia: {md_path} no existe. Se omite.")
            continue
        # Ensure parent directory exists for outputs
        pdf_parent = pdf_path.parent
        pdf_parent.mkdir(parents=True, exist_ok=True)
        if not args.only_pdf:
            docx_parent = docx_path.parent
            docx_parent.mkdir(parents=True, exist_ok=True)
            md_to_docx(md_path, docx_path)
        md_to_pdf(md_path, pdf_path)
        generated.append((md_path, docx_path, pdf_path))

    if generated:
        print("Generados:")
        for _, docx_path, pdf_path in generated:
            if not args.only_pdf:
                print(docx_path)
            print(pdf_path)
    else:
        print("No se generó ningún fichero.")


if __name__ == "__main__":
    main()
